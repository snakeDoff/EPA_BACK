from __future__ import annotations

import re
from collections import defaultdict
from pathlib import Path
from typing import Any

from docx import Document

from app.utils.fio import normalize_fio


EMAIL_RE = re.compile(r"([A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,})")


def _extract_fio_before_email(text: str) -> str:
    """
    Берет все, что стоит перед email, и пытается трактовать это как ФИО.
    Примеры:
    - Иванов Иван Иванович <ivanov@hse.ru>
    - Иванов Иван Иванович ivanov@hse.ru
    """
    match = EMAIL_RE.search(text)
    if not match:
        return ""

    fio_part = text[:match.start()]
    fio_part = (
        fio_part.replace("<", " ")
        .replace(">", " ")
        .replace("(", " ")
        .replace(")", " ")
        .replace(":", " ")
        .replace(";", " ")
        .replace(",", " ")
    )
    fio_part = re.sub(r"\s+", " ", fio_part).strip()

    return normalize_fio(fio_part)


def parse_student_emails_docx(file_path: str | Path) -> dict[str, Any]:
    """
    Возвращает:
    {
        "fio_to_emails": {
            "иванов иван иванович": ["ivanov@hse.ru"]
        },
        "errors": [...]
    }
    """
    doc = Document(file_path)

    text_parts: list[str] = []

    # Параграфы
    for paragraph in doc.paragraphs:
        if paragraph.text and paragraph.text.strip():
            text_parts.append(paragraph.text)

    # Таблицы тоже читаем, потому что в word-выгрузках данные часто лежат в них
    for table in doc.tables:
        for row in table.rows:
            row_text = " ".join(cell.text.strip() for cell in row.cells if cell.text and cell.text.strip())
            if row_text:
                text_parts.append(row_text)

    raw_text = "\n".join(text_parts)

    # Разбиваем текст на куски по строкам и ;
    chunks = re.split(r"[;\n]+", raw_text)

    fio_to_emails: dict[str, list[str]] = defaultdict(list)
    errors: list[dict[str, Any]] = []

    for chunk in chunks:
        chunk = re.sub(r"\s+", " ", chunk).strip()
        if not chunk:
            continue

        email_match = EMAIL_RE.search(chunk)
        if not email_match:
            continue

        email = email_match.group(1).strip().lower()
        fio = _extract_fio_before_email(chunk)

        if not fio:
            errors.append(
                {
                    "segment": chunk,
                    "error": "Не удалось извлечь ФИО из сегмента с email",
                }
            )
            continue

        if email not in fio_to_emails[fio]:
            fio_to_emails[fio].append(email)

    return {
        "fio_to_emails": dict(fio_to_emails),
        "errors": errors,
    }